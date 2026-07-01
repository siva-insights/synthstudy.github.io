# Imports
import os
import re
import sys
import uuid
import base64
import random
import shutil
import subprocess
import requests
import pandas as pd
from datasets import load_dataset
from docx import Document
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from pathlib import Path
from datetime import datetime
from threading import Thread, Lock
import uvicorn
import json
import time
from typing import Literal, Optional

# Constants
OLLAMA_URL = "http://localhost:11434"

# Hidden dir for history and temp CSV; final XLSX always goes to Downloads
_APP_DIR = Path.home() / ".sedg_helper"
_APP_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR = _APP_DIR / "outputs"   # temp CSV lives here
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
DOWNLOADS_DIR = Path.home() / "Downloads"
DOWNLOADS_DIR.mkdir(parents=True, exist_ok=True)
HISTORY_FILE = _APP_DIR / "generation_history.json"

# FastAPI app setup + CORS
app = FastAPI(title="OLSEDG Helper")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/outputs", StaticFiles(directory=OUTPUT_DIR), name="outputs")

# Job state: in-memory dict keyed by job_id and a thread lock for safe updates
JOBS = {}
JOBS_LOCK = Lock()

# Utility functions: context-window estimation, history persistence, timing averages
def estimate_num_ctx(prompt: str) -> int:
    prompt_words = len(str(prompt).split())
    num_ctx = (prompt_words + 500 + 3061) * 2
    return int(num_ctx)
    
def load_history():
    if not HISTORY_FILE.exists():
        return []

    try:
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []


def save_history_entry(entry):
    history = load_history()
    history.append(entry)

    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, indent=2)


def get_average_seconds_per_respondent(model_name):
    history = load_history()
    model_rows = [h for h in history if h.get("model_name") == model_name]

    if not model_rows:
        return None

    values = [h["seconds_taken"] for h in model_rows if h.get("seconds_taken")]
    if not values:
        return None

    return sum(values) / len(values)
    
def update_job(job_id, **kwargs):
    with JOBS_LOCK:
        JOBS[job_id].update(kwargs)


# Data models: Pydantic request/response schemas
class Condition(BaseModel):
    condition_number: int
    condition_name: Optional[str] = None
    stimuli: str


class Question(BaseModel):
    question_number: int
    question_text: str
    scale_points: list[str]
    scale_start: int = 1
    scale_type: str = "discrete"
    max_words: int = 0


class PersonaRecord(BaseModel):
    pid: str
    persona: str


class GenerateRequest(BaseModel):
    study_name: str
    model_provider: Literal["local", "openai", "gemini", "anthropic"] = "local"
    model_name: str
    temperature: float
    sample_count_per_condition: int
    conditions: list[Condition]
    questions: list[Question]
    generic_instruction: Optional[str] = None
    persona_source: Literal["default", "custom", "none"] = "default"
    persona_order: Literal["random", "sequential"] = "random"
    stimuli_assignment: Literal["random", "sequential"] = "random"
    custom_personas: Optional[list[PersonaRecord]] = None

class SaveFileRequest(BaseModel):
    filename: str
    content_base64: str


# API endpoints
@app.get("/health")
def health():
    return {"helper_running": True, "message": "OLSEDG Helper is running"}

@app.post("/save-file")
def save_file(req: SaveFileRequest):
    safe_name = Path(req.filename).name  # strip any path traversal
    dest = OUTPUT_DIR / safe_name
    with open(dest, "wb") as f:
        f.write(base64.b64decode(req.content_base64))
    return {"success": True, "path": str(dest)}

@app.get("/estimate-time/{model_name}/{total_respondents}")
def estimate_time(model_name: str, total_respondents: int):
    avg_seconds = get_average_seconds_per_respondent(model_name)

    if avg_seconds is None:
        return {
            "success": True,
            "has_history": False,
            "message": "No prior timing history available for this model."
        }

    estimated_seconds = avg_seconds * total_respondents
    estimated_minutes = round(estimated_seconds / 60, 1)

    return {
        "success": True,
        "has_history": True,
        "model_name": model_name,
        "average_seconds_per_respondent": round(avg_seconds, 2),
        "total_respondents": total_respondents,
        "estimated_seconds": round(estimated_seconds, 1),
        "estimated_minutes": estimated_minutes,
        "message": f"Estimated generation time: approximately {estimated_minutes} minutes."
    }

@app.get("/check-model/{model_name}")
def check_model(model_name: str):
    try:
        response = requests.get(f"{OLLAMA_URL}/api/tags", timeout=5)
        response.raise_for_status()

        models = response.json().get("models", [])
        installed_models = [m["name"] for m in models]
        model_installed = model_name in installed_models

        return {
            "helper_running": True,
            "ollama_running": True,
            "model": model_name,
            "model_installed": model_installed,
            "ready_to_use": model_installed,
            "installed_models": installed_models,
        }

    except requests.exceptions.RequestException:
        return {
            "helper_running": True,
            "ollama_running": False,
            "model_installed": False,
            "ready_to_use": False,
            "message": "Ollama is not running.",
        }


def sample_personas(df_small: pd.DataFrame, total_needed: int, sequential: bool = False):
    if sequential:
        # Repeat the list cyclically until total_needed rows are available, then take first N
        reps = (total_needed // len(df_small)) + 1
        return pd.concat([df_small] * reps, ignore_index=True).iloc[:total_needed].reset_index(drop=True)
    replace = total_needed > len(df_small)
    return df_small.sample(
        n=total_needed,
        replace=replace,
        random_state=random.randint(1, 999999)
    ).reset_index(drop=True)


def load_personas(
    total_needed: int,
    custom_personas: Optional[list[PersonaRecord]] = None,
    use_personas: bool = True,
    sequential: bool = False
):
    if not use_personas:
        return pd.DataFrame([
            {"pid": f"no_persona_{i + 1}", "persona_summary": ""}
            for i in range(total_needed)
        ])

    if custom_personas:
        df_small = pd.DataFrame([
            {"pid": p.pid, "persona_summary": p.persona}
            for p in custom_personas
        ]).dropna()

        df_small["persona_summary"] = df_small["persona_summary"].astype(str).str.strip()
        df_small = df_small[df_small["persona_summary"] != ""]

        if df_small.empty:
            raise ValueError("Custom persona file must include at least one non-empty persona.")

        if sequential:
            df_small = df_small.sort_values("pid").reset_index(drop=True)

        return sample_personas(df_small, total_needed, sequential=sequential)

    ds = load_dataset("LLM-Digital-Twin/Twin-2K-500", "full_persona")

    if hasattr(ds, "keys"):
        split_name = list(ds.keys())[0]
        df = ds[split_name].to_pandas()
    else:
        df = ds.to_pandas()

    df_small = df[["pid", "persona_summary"]].dropna().copy()
    if sequential:
        df_small = df_small.sort_values("pid").reset_index(drop=True)
    return sample_personas(df_small, total_needed, sequential=sequential)


def build_prompt(
    persona,
    stimuli,
    questions,
    generic_instruction: Optional[str] = None,
    include_persona: bool = True
):
    question_lookup = {}

    for q in questions:
        if q.scale_type == "text":
            max_words_str = f"{q.max_words} words" if q.max_words > 0 else "unlimited"
            embedded_question = f"""
[Q{q.question_number}]
Question: {q.question_text}
Response type: Text
Max words: {max_words_str}
Please provide a text response within the allowed word count.
[/Q{q.question_number}]
""".strip()
        else:
            min_code = q.scale_start
            max_code = q.scale_start + len(q.scale_points) - 1

            scale_text = "\n".join(
                [f"{q.scale_start + i} = {label}" for i, label in enumerate(q.scale_points)]
            )

            if q.scale_type == "continuous":
                range_instruction = (
                    f"Allowed response range: {min_code} to {max_code}\n"
                    f"Please select a number in the range."
                )
            else:
                range_instruction = (
                    f"Allowed response codes: {min_code} to {max_code}\n"
                    f"Please select an integer in the range."
                )

            response_type_label = "Continuous" if q.scale_type == "continuous" else "Discrete"
            embedded_question = f"""
[Q{q.question_number}]
Question: {q.question_text}
Response type: {response_type_label}
{range_instruction}
Response scale:
{scale_text}
[/Q{q.question_number}]
""".strip()

        question_lookup[q.question_number] = embedded_question

    embedded_stimuli = stimuli

    for q in questions:
        placeholder = f"{{Q{q.question_number}}}"
        embedded_stimuli = embedded_stimuli.replace(
            placeholder,
            question_lookup[q.question_number]
        )

    answer_template = "\n".join(
        [f'Q{q.question_number}="?"' if q.scale_type == "text" else f"Q{q.question_number}=?"
         for q in questions]
    )

    default_prompt_template = """
You are simulating one synthetic survey respondent.

Your task:
1. Read the respondent persona.
2. Read the study materials exactly as a survey participant would see them.
3. Answer each embedded survey question from this respondent's perspective.
4. Use the respondent persona, the study materials, response scale, and scale type for each question when choosing answers.

Respondent persona:
{persona}

Study materials with embedded questions:
{embedded_stimuli}

Important rules:
* Use only the allowed response codes or response format for each question.
* Match each answer to the question type:
    * Discrete scale: use an integer.
    * Continuous scale: use a number/float.
    * Text response: write text within the allowed word limit.
* Do not use values below the minimum code or above the maximum code.
* Return only the final responses. Do not include explanations, markdown, or repeat questions.
* Return answers only in this format: {answer_template}
""".strip()

    no_persona_prompt_template = """
You are simulating one synthetic survey respondent.

Your task:
1. Read the study materials exactly as a survey participant would see them.
2. Answer each embedded survey question from a participant's perspective.
3. Use the respondent persona, the study materials, response scale, and scale type for each question when choosing answers.

Study materials with embedded questions:
{embedded_stimuli}

Important rules:
* Use only the allowed response codes or response format for each question.
* Match each answer to the question type:
    * Discrete scale: use an integer.
    * Continuous scale: use a number/float.
    * Text response: write text within the allowed word limit.
* Do not use values below the minimum code or above the maximum code.
* Return only the final responses. Do not include explanations, markdown, or repeat questions.
* Return answers only in this format: {answer_template}
""".strip()

    prompt_template = (generic_instruction or "").strip()

    if not prompt_template:
        prompt_template = default_prompt_template if include_persona else no_persona_prompt_template

    if not include_persona:
        prompt_template = re.sub(
            r"\n?Respondent persona:\s*\n\{persona\}\s*\n",
            "\n",
            prompt_template
        )
        prompt_template = re.sub(
            r"\n?\d+\.\s*Read the respondent persona\.\s*\n?",
            "\n",
            prompt_template
        )
        prompt_template = re.sub(r"\n2\. ", "\n1. ", prompt_template)
        prompt_template = re.sub(r"\n3\. ", "\n2. ", prompt_template)
        prompt_template = re.sub(r"\n4\. ", "\n3. ", prompt_template)
        prompt_template = prompt_template.replace("this respondent's perspective", "a participant's perspective")

    prompt = prompt_template.replace("{persona}", str(persona) if include_persona else "")
    prompt = prompt.replace("{embedded_stimuli}", str(embedded_stimuli))
    prompt = prompt.replace("{answer_template}", str(answer_template))

    return prompt


def call_ollama(model_name, prompt, temperature):
    num_ctx = estimate_num_ctx(prompt)

    response = requests.post(
        f"{OLLAMA_URL}/api/generate",
        json={
            "model": model_name,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": temperature,
                "num_ctx": num_ctx
            }
        },
        timeout=900
    )

    response.raise_for_status()
    return response.json().get("response", "")


def parse_answers(raw_text, questions):
    answers = {}
    invalid_questions = []

    for q in questions:
        q_num = q.question_number

        if q.scale_type == "text":
            match = re.search(rf'Q{q_num}\s*=\s*"([^"]*)"', raw_text, re.IGNORECASE)
            if not match:
                match = re.search(rf"Q{q_num}\s*=\s*(.+)", raw_text, re.IGNORECASE)
            if match:
                answers[f"Q{q_num}"] = match.group(1).strip()
            else:
                answers[f"Q{q_num}"] = ""
                invalid_questions.append(f"Q{q_num}")
        else:
            min_code = q.scale_start
            max_code = q.scale_start + len(q.scale_points) - 1
            is_continuous = q.scale_type == "continuous"

            pattern = rf"Q{q_num}\s*=\s*(-?\d+\.?\d*)"
            match = re.search(pattern, raw_text, re.IGNORECASE)

            if match:
                value = float(match.group(1)) if is_continuous else int(match.group(1))
                answers[f"Q{q_num}"] = value

                if not (min_code <= value <= max_code):
                    invalid_questions.append(f"Q{q_num}")
            else:
                answers[f"Q{q_num}"] = ""
                invalid_questions.append(f"Q{q_num}")

    validation = "valid" if not invalid_questions else "invalid"

    return answers, validation, invalid_questions
    
def get_valid_response_with_retries(model_name, prompt, temperature, questions, max_retries=5):
    last_raw_response = ""
    last_answers = {}
    last_validation = "invalid"
    last_invalid_questions = []

    for attempt in range(1, max_retries + 1):
        raw_response = call_ollama(model_name, prompt, temperature)
        answers, validation, invalid_questions = parse_answers(raw_response, questions)

        if validation == "valid":
            return answers, validation, invalid_questions, attempt, raw_response

        last_raw_response = raw_response
        last_answers = answers
        last_validation = validation
        last_invalid_questions = invalid_questions

    return last_answers, last_validation, last_invalid_questions, max_retries, last_raw_response

def create_docx(data, filepath):
    doc = Document()

    doc.add_heading("OLSEDG Study Inputs", level=1)

    doc.add_heading("Study Information", level=2)
    doc.add_paragraph(f"Study Name: {data.study_name}")
    doc.add_paragraph(f"Model: {data.model_name}")
    doc.add_paragraph(f"Temperature: {data.temperature}")
    doc.add_paragraph(f"Samples per Condition: {data.sample_count_per_condition}")
    doc.add_paragraph(f"Number of Conditions: {len(data.conditions)}")

    doc.add_heading("Conditions and Stimuli", level=2)
    for c in data.conditions:
        doc.add_heading(f"Condition {c.condition_number}", level=3)
        doc.add_paragraph(c.stimuli)

    doc.add_heading("Questions and Response Scales", level=2)
    for q in data.questions:
        doc.add_heading(f"Q{q.question_number}: {q.question_text}", level=3)
        for i, label in enumerate(q.scale_points, start=q.scale_start):
            doc.add_paragraph(f"{i} = {label}")

    doc.save(filepath)


# Generation job worker: runs in a background thread, generates one respondent per iteration
def run_generation_job(job_id: str, data: GenerateRequest):
    try:
        total_needed = data.sample_count_per_condition * len(data.conditions)

        update_job(
            job_id,
            status="loading_personas",
            message="Loading respondent personas...",
            completed=0,
            total=total_needed
        )

        use_personas = data.persona_source != "none"
        custom_personas = data.custom_personas if data.persona_source == "custom" else None
        sequential_personas = getattr(data, "persona_order", "random") == "sequential"
        df_personas = load_personas(total_needed, custom_personas, use_personas, sequential=sequential_personas)

        condition_numbers = []
        for c in data.conditions:
            condition_numbers.extend([c.condition_number] * data.sample_count_per_condition)

        if getattr(data, "stimuli_assignment", "random") == "random":
            random.shuffle(condition_numbers)
        # sequential: already in order (C1 × N, C2 × N, …) — no shuffle needed

        condition_lookup = {c.condition_number: c.stimuli for c in data.conditions}
        condition_name_lookup = {
            c.condition_number: c.condition_name or f"Condition {c.condition_number}"
            for c in data.conditions
        }

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_model_name = data.model_name.replace(":", "-").replace("/", "-")
        num_conditions = len(data.conditions)
        samples_per_condition = data.sample_count_per_condition

        xlsx_filename = f"SEDG_{total_needed}samples_{timestamp}.xlsx"
        csv_path = OUTPUT_DIR / f"_tmp_{timestamp}.csv"   # temp; converted to xlsx at end

        question_columns = [f"Q{q.question_number}" for q in data.questions]
        columns = [
            "respondent_id",
            "study_name",
            "pid",
            "persona_summary",
            "condition",
            "condition_name",
            "condition_stimuli",
            "model_name",
            "temperature",
            "validation",
            "invalid_questions",
            "retry_count",
            "seconds_taken",
            "prompt",
            "prompt_words",
            "num_ctx_used",
            "raw_response"
        ] + question_columns
        pd.DataFrame(columns=columns).to_csv(csv_path, index=False)

        update_job(
            job_id,
            status="generating",
            message=f"Generating synthetic responses. Output is being saved continuously to {csv_path}",
            completed=0,
            total=total_needed,
            percent=0,
        )

        job_start_time = time.time()

        for i in range(total_needed):
            if JOBS.get(job_id, {}).get("stop_requested"):
                if csv_path.exists():
                    _xp = DOWNLOADS_DIR / xlsx_filename
                    pd.read_csv(csv_path).to_excel(_xp, index=False, engine="openpyxl")
                    csv_path.unlink(missing_ok=True)
                    update_job(job_id, status="stopped", xlsx_path=str(_xp),
                               message=f"Generation stopped after {i} respondents.")
                else:
                    update_job(job_id, status="stopped",
                               message=f"Generation stopped after {i} respondents.")
                return

            respondent_id = i + 1
            condition_number = condition_numbers[i]
            stimuli = condition_lookup[condition_number]

            # Only the questions actually embedded ({Qn} present) in this condition's
            # stimuli are shown to the model, so only those should be expected back.
            embedded_questions = [
                q for q in data.questions if f"{{Q{q.question_number}}}" in stimuli
            ]

            pid = df_personas.loc[i, "pid"]
            persona = df_personas.loc[i, "persona_summary"]

            prompt = build_prompt(
                persona,
                stimuli,
                embedded_questions,
                data.generic_instruction,
                include_persona=use_personas
            )

            start_time = time.time()

            answers, validation, invalid_questions, retry_count, raw_response = get_valid_response_with_retries(
                data.model_name,
                prompt,
                data.temperature,
                embedded_questions,
                max_retries=5
            )

            end_time = time.time()
            seconds_taken = round(end_time - start_time, 3)

            save_history_entry({
                "timestamp": datetime.now().isoformat(),
                "model_name": data.model_name,
                "seconds_taken": seconds_taken,
                "respondent_id": respondent_id,
                "condition": condition_number
            })

            prompt_words = len(str(prompt).split())
            num_ctx_used = estimate_num_ctx(prompt)
                        
            row = {
                "respondent_id": respondent_id,
                "study_name": data.study_name,
                "pid": pid,
                "persona_summary": persona,
                "condition": condition_number,
                "condition_name": condition_name_lookup[condition_number],
                "condition_stimuli": stimuli,
                "model_name": data.model_name,
                "temperature": data.temperature,
                "validation": validation,
                "invalid_questions": ",".join(invalid_questions),
                "retry_count": retry_count,
                "seconds_taken": seconds_taken,
                "prompt": prompt,
                "prompt_words": prompt_words,
                "num_ctx_used": num_ctx_used,
                "raw_response": str(raw_response).replace("\n", " | "),
            }
            
            # Safety check: if answers accidentally comes as a tuple, take only the answers dictionary
            if isinstance(answers, tuple):
                answers = answers[0]
            
            # Safety check: if answers is still not a dictionary, stop with a clear error
            if not isinstance(answers, dict):
                raise ValueError(f"answers should be a dictionary, but got {type(answers)}: {answers}")
            
            # Add each question answer manually instead of using row.update()
            for q in data.questions:
                q_col = f"Q{q.question_number}"
                row[q_col] = answers.get(q_col, "")
            pd.DataFrame([row]).to_csv(
                csv_path,
                mode="a",
                header=False,
                index=False
            )

            completed = i + 1
            pending = total_needed - completed
            percent = round((completed / total_needed) * 100, 1)

            elapsed_seconds = time.time() - job_start_time
            avg_seconds_current_run = elapsed_seconds / completed
            estimated_remaining_seconds = avg_seconds_current_run * pending
            estimated_total_seconds = elapsed_seconds + estimated_remaining_seconds

            update_job(
                job_id,
                completed=completed,
                total=total_needed,
                pending=pending,
                percent=percent,
                elapsed_seconds=round(elapsed_seconds, 1),
                average_seconds_per_respondent=round(avg_seconds_current_run, 2),
                estimated_remaining_seconds=round(estimated_remaining_seconds, 1),
                estimated_total_seconds=round(estimated_total_seconds, 1),
                message=f"{completed}/{total_needed} respondents completed. {pending} remaining."
            )

        # Convert the temp CSV to XLSX then remove the CSV
        xlsx_path = DOWNLOADS_DIR / xlsx_filename
        pd.read_csv(csv_path).to_excel(xlsx_path, index=False, engine="openpyxl")
        csv_path.unlink(missing_ok=True)

        update_job(
            job_id,
            status="complete",
            message="Generation complete",
            completed=total_needed,
            total=total_needed,
            pending=0,
            percent=100,
            estimated_remaining_seconds=0,
            xlsx_path=str(xlsx_path),
        )

    except Exception as e:
        update_job(
            job_id,
            status="error",
            message=str(e)
        )
        

@app.get("/preview-personas/{sample_count}")
def preview_personas(sample_count: int):
    if sample_count < 1:
        sample_count = 1

    df_personas = load_personas(sample_count, use_personas=True)

    return {
        "success": True,
        "personas": [
            {
                "pid": str(row["pid"]),
                "persona": str(row["persona_summary"])
            }
            for _, row in df_personas.iterrows()
        ]
    }

@app.post("/generate")
def generate(data: GenerateRequest):
    total_needed = data.sample_count_per_condition * len(data.conditions)
    job_id = str(uuid.uuid4())

    JOBS[job_id] = {
        "job_id": job_id,
        "status": "started",
        "message": "Generation started",
        "completed": 0,
        "total": total_needed,
    }

    thread = Thread(target=run_generation_job, args=(job_id, data), daemon=True)
    thread.start()

    return {
        "success": True,
        "job_id": job_id,
        "total_responses": total_needed
    }


@app.get("/progress/{job_id}")
def get_progress(job_id: str):
    if job_id not in JOBS:
        return {
            "success": False,
            "message": "Job not found"
        }

    return {
        "success": True,
        **JOBS[job_id]
    }

@app.post("/stop/{job_id}")
def stop_job(job_id: str):
    if job_id not in JOBS:
        return {"success": False, "message": "Job not found"}
    with JOBS_LOCK:
        JOBS[job_id]["stop_requested"] = True
    return {"success": True}

@app.get("/download-result/{job_id}")
def download_result(job_id: str):
    from fastapi.responses import FileResponse
    xlsx_path = JOBS.get(job_id, {}).get("xlsx_path")
    if not xlsx_path or not Path(xlsx_path).exists():
        return {"success": False, "message": "File not ready"}
    return FileResponse(
        path=xlsx_path,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename=Path(xlsx_path).name
    )

# tkinter GUI entry point
if __name__ == "__main__":
    import tkinter as tk
    from tkinter import ttk, messagebox
    import webbrowser

    PORT = 8000
    OLLAMA_BASE = "http://localhost:11434"

    # Color palette
    C = {
        "bg":      "#f1f5f9",   # page background
        "card":    "#ffffff",   # card / tab content
        "blue":    "#1e40af",   # primary / header
        "blue2":   "#1d4ed8",   # button hover
        "sky":     "#93c5fd",   # header muted text
        "green":   "#16a34a",   # success
        "green_bg":"#f0fdf4",   # success badge bg
        "red":     "#dc2626",   # error
        "red_bg":  "#fef2f2",   # error badge bg
        "text":    "#1e293b",   # primary text
        "muted":   "#64748b",   # secondary text
        "border":  "#e2e8f0",   # subtle border
        "row":     "#f8fafc",   # alternate row bg
    }

    # Model list / sizes
    MODEL_SIZES = {
        "deepseek-r1:14b": "~9.0 GB", "gemma3:4b": "~3.3 GB",  "gemma3:12b": "~8.1 GB",
        "llama3.1:8b":     "~4.7 GB", "llama3.2:3b": "~2.0 GB", "mistral:7b-instruct": "~4.1 GB",
        "mistral-small3.2:24b": "~14 GB", "qwen2.5:7b": "~4.4 GB",
        "qwen2.5:7b-instruct":  "~4.4 GB", "qwen2.5:14b": "~9.0 GB", "qwen3:14b": "~9.3 GB",
    }
    MODEL_LIST = list(MODEL_SIZES.keys()) + ["Other Model"]

    # Ollama helpers: detect, start, stop, and list installed models
    def _ollama_ok():
        try:
            return requests.get(f"{OLLAMA_BASE}/api/tags", timeout=3).status_code == 200
        except Exception:
            return False

    def _ollama_installed():
        return os.path.exists("/Applications/Ollama.app") or shutil.which("ollama") is not None

    def _start_ollama():
        if os.path.exists("/Applications/Ollama.app"):
            subprocess.Popen(["open", "-a", "Ollama"],
                             stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        else:
            cli = shutil.which("ollama")
            if cli:
                subprocess.Popen([cli, "serve"],
                                 stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

    def _ensure_ollama_running():
        """Start Ollama if needed. Returns True when running."""
        if _ollama_ok():
            return True
        if not _ollama_installed():
            return False
        _start_ollama()
        for _ in range(12):
            time.sleep(1)
            if _ollama_ok():
                return True
        return False

    def _quit_ollama():
        # Gracefully quit the macOS app bundle (menu bar icon)
        subprocess.run(["osascript", "-e", 'quit app "Ollama"'],
                       stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        # Force-kill both the app bundle process (Ollama) and the background
        # server process (ollama) — covers both .app and CLI installs
        subprocess.run(["killall", "Ollama"],
                       stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        subprocess.run(["killall", "ollama"],
                       stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

    def _installed_models():
        try:
            r = requests.get(f"{OLLAMA_BASE}/api/tags", timeout=3)
            return r.json().get("models", []) if r.ok else []
        except Exception:
            return []

    # FastAPI background server thread: starts uvicorn on localhost:8000
    config = uvicorn.Config(app, host="127.0.0.1", port=PORT, log_level="error")
    server = uvicorn.Server(config)
    Thread(target=server.run, daemon=True).start()

    # Window + header setup: root window, header bar, notebook widget and shared button helpers
    root = tk.Tk()
    root.title("OLSEDG Helper")
    root.geometry("500x460")
    root.resizable(True, True)
    root.configure(bg=C["bg"])

    # Header bar
    hdr = tk.Frame(root, bg=C["blue"], height=46)
    hdr.pack(fill="x"); hdr.pack_propagate(False)
    tk.Label(hdr, text="OLSEDG Helper", bg=C["blue"], fg="white",
             font=("Helvetica", 14, "bold")).pack(side="left", padx=16)

    # Notebook styles
    sty = ttk.Style(); sty.theme_use("clam")
    sty.configure("TNotebook", background=C["bg"], borderwidth=0, tabmargins=0)
    sty.configure("TNotebook.Tab", padding=[16, 6], font=("Helvetica", 10),
                  background=C["border"], foreground=C["muted"])
    sty.map("TNotebook.Tab",
            background=[("selected", C["card"])],
            foreground=[("selected", C["blue"])])
    sty.configure("TFrame", background=C["card"])
    sty.configure("Horizontal.TProgressbar",
                  troughcolor=C["border"], background=C["blue"],
                  borderwidth=0, thickness=8)

    nb = ttk.Notebook(root)
    nb.pack(fill="both", expand=True, padx=10, pady=(6, 10))

    def _btn(parent, text, cmd, **kw):
        """Flat blue button matching SEDG style."""
        b = tk.Button(parent, text=text, command=cmd,
                      bg=C["blue"], fg="white", activebackground=C["blue2"],
                      activeforeground="white", relief="flat", bd=0,
                      padx=14, pady=7, font=("Helvetica", 10, "bold"),
                      cursor="hand2", **kw)
        return b

    def _ghost_btn(parent, text, cmd):
        """Muted secondary button."""
        return tk.Button(parent, text=text, command=cmd,
                         bg=C["border"], fg=C["text"], activebackground="#cbd5e1",
                         relief="flat", bd=0, padx=10, pady=5,
                         font=("Helvetica", 9), cursor="hand2")

    def _card(parent):
        """White card frame with a subtle border feel."""
        return tk.Frame(parent, bg=C["card"], padx=20, pady=18)

    # Tab 1 (Ollama Setup): validate Ollama installation and auto-start if needed
    tab1 = ttk.Frame(nb); nb.add(tab1, text="1. Ollama Setup")
    p1 = _card(tab1); p1.pack(fill="both", expand=True)

    # Title + tooltip
    th = tk.Frame(p1, bg=C["card"]); th.pack(anchor="w")
    tk.Label(th, text="Validate Ollama Setup", font=("Helvetica", 13, "bold"),
             bg=C["card"], fg=C["text"]).pack(side="left")
    tip = tk.Label(th, text="  ⓘ", font=("Helvetica", 11), fg=C["muted"],
                   bg=C["card"], cursor="hand2")
    tip.pack(side="left")

    _tw = [None]
    def _show_tip(e):
        if _tw[0]: return
        w = tk.Toplevel(root); w.overrideredirect(True)
        w.configure(bg=C["text"])
        tk.Label(w, text=("Ollama runs open-source LLMs locally on your computer.\n"
                           "OLSEDG Helper sends prompts to it to generate synthetic\n"
                           "responses — no data ever leaves your machine."),
                 bg=C["text"], fg="white", font=("Helvetica", 10),
                 padx=12, pady=10, justify="left").pack()
        w.geometry(f"+{e.x_root+8}+{e.y_root+8}"); _tw[0] = w
    def _hide_tip(e):
        if _tw[0]: _tw[0].destroy(); _tw[0] = None
    tip.bind("<Enter>", _show_tip); tip.bind("<Leave>", _hide_tip)

    tk.Frame(p1, bg=C["border"], height=1).pack(fill="x", pady=(10, 14))

    t1_stat = tk.Label(p1, text="", font=("Helvetica", 12), bg=C["card"])
    t1_stat.pack(anchor="w")
    t1_det  = tk.Label(p1, text="", font=("Helvetica", 10), fg=C["muted"],
                        bg=C["card"], wraplength=440)
    t1_det.pack(anchor="w", pady=(4, 0))
    t1_lnk  = tk.Label(p1, text="", font=("Helvetica", 10, "underline"),
                        fg=C["blue"], bg=C["card"], cursor="hand2")
    t1_lnk.pack(anchor="w", pady=(4, 0))
    t1_lnk.bind("<Button-1>", lambda e: webbrowser.open("https://ollama.ai"))

    tk.Frame(p1, bg=C["card"], height=6).pack()
    t1_btn = tk.Button(p1, text="Check Again",
                       command=lambda: Thread(target=do_check, daemon=True).start(),
                       bg=C["red_bg"], fg=C["red"], activebackground="#fecaca",
                       relief="flat", bd=0, padx=14, pady=7,
                       font=("Helvetica", 10, "bold"), cursor="hand2")
    t1_btn.pack(anchor="w")

    def do_check():
        root.after(0, lambda: (t1_stat.config(text="Checking…", fg=C["muted"]),
                               t1_det.config(text=""), t1_lnk.config(text="")))
        ok = _ollama_ok()
        if not ok and _ollama_installed():
            root.after(0, lambda: t1_stat.config(text="Ollama found — starting…", fg=C["muted"]))
            _start_ollama()
            for _ in range(12):
                time.sleep(1)
                if _ollama_ok(): ok = True; break
        def u():
            if ok:
                t1_stat.config(text="✓  Ollama is running.", fg=C["green"])
                t1_det.config(text="Proceed to Install Model to choose your LLM.")
                t1_lnk.config(text="")
            elif _ollama_installed():
                t1_stat.config(text="✗  Could not start Ollama.", fg=C["red"])
                t1_det.config(text="Try opening Ollama manually, then click Check Again.")
                t1_lnk.config(text="")
            else:
                t1_stat.config(text="✗  Ollama not installed.", fg=C["red"])
                t1_det.config(text="Download and install Ollama, then click Check Again.")
                t1_lnk.config(text="↗  ollama.ai — download Ollama")
        root.after(0, u)

    # Tab 2 (Install Model): model picker, download progress bar, and ollama pull integration
    tab2 = ttk.Frame(nb); nb.add(tab2, text="2. Install Model")
    p2 = _card(tab2); p2.pack(fill="both", expand=True)

    tk.Label(p2, text="Install a Model", font=("Helvetica", 13, "bold"),
             bg=C["card"], fg=C["text"]).pack(anchor="w")
    tk.Label(p2, text="Choose a model for generating synthetic survey responses.",
             font=("Helvetica", 10), fg=C["muted"], bg=C["card"]).pack(anchor="w", pady=(2, 0))
    tk.Frame(p2, bg=C["border"], height=1).pack(fill="x", pady=(10, 12))

    t2_var = tk.StringVar(value="Select a model")
    t2_cb  = ttk.Combobox(p2, textvariable=t2_var, values=MODEL_LIST,
                           state="readonly", width=32, font=("Helvetica", 11))
    t2_cb.pack(anchor="w")

    # "Other" entry (shown only when Other is selected)
    t2_of = tk.Frame(p2, bg=C["card"])
    t2_ov = tk.StringVar()
    tk.Entry(t2_of, textvariable=t2_ov, font=("Helvetica", 11), width=34,
             bg=C["row"], fg=C["text"], insertbackground=C["text"],
             relief="flat", bd=1).pack(anchor="w", pady=(6, 0))
    tk.Label(t2_of, text="Enter model name (e.g. llama3.2:1b from ollama.com/library)",
             font=("Helvetica", 9), fg=C["muted"], bg=C["card"]).pack(anchor="w", pady=(3, 0))

    # Info badge (size + installed status)
    t2_badge = tk.Frame(p2, bg=C["card"]); t2_badge.pack(anchor="w", pady=(10, 0))
    t2_badge_lbl = tk.Label(t2_badge, text="", font=("Helvetica", 10), bg=C["card"])
    t2_badge_lbl.pack(anchor="w")

    # Progress frame
    t2_pf = tk.Frame(p2, bg=C["card"])
    t2_pg = ttk.Progressbar(t2_pf, style="Horizontal.TProgressbar",
                             mode="determinate", length=440)
    t2_pg.pack(anchor="w")
    t2_pl = tk.Label(t2_pf, text="", font=("Helvetica", 9), fg=C["muted"], bg=C["card"])
    t2_pl.pack(anchor="w", pady=(3, 0))

    t2_stat = tk.Label(p2, text="", font=("Helvetica", 10), bg=C["card"], wraplength=440)
    t2_stat.pack(anchor="w", pady=(6, 0))
    t2_lnk  = tk.Label(p2, text="", font=("Helvetica", 10, "underline"),
                        fg=C["blue"], bg=C["card"], cursor="hand2")
    t2_lnk.pack(anchor="w")
    t2_lnk.bind("<Button-1>", lambda e: webbrowser.open("https://synthstudy.vercel.app"))

    tk.Frame(p2, bg=C["card"], height=4).pack()
    t2_btn = tk.Button(p2, text="Install Model", command=lambda: None,
                       bg=C["red_bg"], fg=C["red"], activebackground="#fecaca",
                       relief="flat", bd=0, padx=14, pady=7,
                       font=("Helvetica", 10, "bold"), cursor="hand2")
    t2_btn.pack(anchor="w")

    def _check_installed_async(model_name, callback):
        def run():
            names = [m.get("name", "") for m in _installed_models()]
            callback(model_name in names)
        Thread(target=run, daemon=True).start()

    def on_t2_select(e=None):
        val = t2_var.get()
        t2_pf.pack_forget()
        t2_stat.config(text=""); t2_lnk.config(text="")
        t2_badge_lbl.config(text="", bg=C["card"]); t2_badge.config(bg=C["card"])
        if val == "Other Model":
            t2_of.pack(anchor="w", pady=(6, 0), before=t2_badge)
            return
        t2_of.pack_forget()
        if val in ("Select a model", ""):
            return
        sz = MODEL_SIZES.get(val, "")
        t2_badge_lbl.config(text=f"Download size: {sz}" if sz else "Size: checking…",
                             fg=C["muted"], bg=C["card"]); t2_badge.config(bg=C["card"])
        def on_installed(already):
            def u():
                if already:
                    t2_badge.config(bg=C["green_bg"])
                    t2_badge_lbl.config(
                        text=f"✓  Already installed" + (f"  ·  {sz}" if sz else ""),
                        fg=C["green"], bg=C["green_bg"])
                    t2_stat.config(
                        text=f"{val} is ready. You can now generate synthetic responses on SEDG.",
                        fg=C["green"])
                    t2_lnk.config(text="↗  Open SEDG at synthstudy.vercel.app")
                    t2_btn.config(text="Reinstall")
                else:
                    t2_badge.config(bg=C["card"])
                    t2_badge_lbl.config(
                        text=f"Download size: {sz}" if sz else "",
                        fg=C["muted"], bg=C["card"])
                    t2_stat.config(text="")
                    t2_lnk.config(text="")
                    t2_btn.config(text="Install Model")
            root.after(0, u)
        _check_installed_async(val, on_installed)

    t2_cb.bind("<<ComboboxSelected>>", on_t2_select)

    def do_install():
        val   = t2_var.get()
        model = (t2_ov.get().strip() if val == "Other Model"
                 else (None if val in ("Select a model", "") else val))
        if not model:
            t2_stat.config(text="Please select or enter a model name.", fg=C["red"]); return

        t2_btn.config(state="disabled", text="Installing…")
        t2_stat.config(text="", fg=C["muted"])
        t2_pf.pack(anchor="w", pady=(10, 0), before=t2_stat)
        t2_pg["value"] = 0; t2_pl.config(text=""); t2_lnk.config(text="")

        def on_p(status, total, completed):
            def u():
                if total and completed:
                    t2_pg["value"] = int(completed / total * 100)
                    t2_pl.config(text=f"{status}  ·  "
                                      f"{completed/1024**2:.0f} / {total/1024**2:.0f} MB")
                else:
                    t2_pl.config(text=status)
            root.after(0, u)

        def on_done():
            def u():
                t2_pg["value"] = 100; t2_pl.config(text="")
                t2_stat.config(text=f"✓  {model} is ready to use.", fg=C["green"])
                t2_lnk.config(text="↗  Open SEDG at synthstudy.vercel.app")
                t2_btn.config(state="normal", text="Install Model")
                t2_badge.config(bg=C["green_bg"])
                t2_badge_lbl.config(text="✓  Installed", fg=C["green"], bg=C["green_bg"])
            root.after(0, u)

        def on_err(msg):
            def u():
                t2_stat.config(text=f"Error: {msg}", fg=C["red"])
                t2_pf.pack_forget()
                t2_btn.config(state="normal", text="Install Model")
            root.after(0, u)

        def run():
            # Ensure Ollama is running before pulling
            root.after(0, lambda: t2_pl.config(text="Checking Ollama…"))
            if not _ensure_ollama_running():
                on_err("Ollama is not running. Validate setup in tab 1 first."); return
            root.after(0, lambda: t2_pl.config(text="Connecting to Ollama…"))
            try:
                resp = requests.post(f"{OLLAMA_BASE}/api/pull",
                                     json={"name": model}, stream=True, timeout=None)
                for line in resp.iter_lines():
                    if line:
                        d = json.loads(line)
                        on_p(d.get("status", ""), d.get("total", 0), d.get("completed", 0))
                        if d.get("status") == "success":
                            on_done(); return
                on_done()
            except Exception as ex:
                on_err(str(ex))

        Thread(target=run, daemon=True).start()

    t2_btn.config(command=do_install)

    # Tab 3 (Manage Models): scrollable list of installed models with per-model uninstall
    tab3 = ttk.Frame(nb); nb.add(tab3, text="3. Manage Models")
    p3 = _card(tab3); p3.pack(fill="both", expand=True)

    t3h = tk.Frame(p3, bg=C["card"]); t3h.pack(fill="x")
    tk.Label(t3h, text="Installed Models", font=("Helvetica", 13, "bold"),
             bg=C["card"], fg=C["text"]).pack(side="left")
    t3_ref = _ghost_btn(t3h, "↻  Refresh", lambda: refresh_t3())
    t3_ref.pack(side="right")

    tk.Frame(p3, bg=C["border"], height=1).pack(fill="x", pady=(10, 10))

    # Scrollable model list
    t3_canvas_frame = tk.Frame(p3, bg=C["card"])
    t3_canvas_frame.pack(fill="both", expand=True)
    t3_canvas = tk.Canvas(t3_canvas_frame, bg=C["card"], highlightthickness=0)
    t3_vsb = ttk.Scrollbar(t3_canvas_frame, orient="vertical", command=t3_canvas.yview)
    t3_canvas.configure(yscrollcommand=t3_vsb.set)
    t3_vsb.pack(side="right", fill="y")
    t3_canvas.pack(side="left", fill="both", expand=True)
    t3_list = tk.Frame(t3_canvas, bg=C["card"])
    t3_list_id = t3_canvas.create_window((0, 0), window=t3_list, anchor="nw")

    def _t3_on_resize(e):
        t3_canvas.itemconfig(t3_list_id, width=e.width)
    t3_canvas.bind("<Configure>", _t3_on_resize)

    def _t3_on_frame_resize(e):
        t3_canvas.configure(scrollregion=t3_canvas.bbox("all"))
    t3_list.bind("<Configure>", _t3_on_frame_resize)

    # Mouse wheel scroll
    def _t3_scroll(e):
        t3_canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")
    t3_canvas.bind_all("<MouseWheel>", _t3_scroll)

    t3_stat = tk.Label(p3, text="", font=("Helvetica", 10), bg=C["card"])
    t3_stat.pack(anchor="w", pady=(6, 0))

    def refresh_t3():
        for w in t3_list.winfo_children(): w.destroy()
        t3_stat.config(text="")
        tk.Label(t3_list, text="Loading…", fg=C["muted"], bg=C["card"],
                 font=("Helvetica", 10)).pack(anchor="w")

        def do():
            models = _installed_models()
            def upd():
                for w in t3_list.winfo_children(): w.destroy()
                if not models:
                    tk.Label(t3_list, text="No models installed.", fg=C["muted"],
                             bg=C["card"], font=("Helvetica", 10)).pack(anchor="w")
                    return
                for m in models:
                    name = m.get("name", "")
                    sz   = m.get("size", 0)
                    row  = tk.Frame(t3_list, bg=C["row"], padx=10, pady=7)
                    row.pack(fill="x", pady=3)
                    tk.Label(row, text=name, font=("Helvetica", 10, "bold"),
                             bg=C["row"], fg=C["text"]).pack(side="left")
                    if sz:
                        tk.Label(row, text=f"  {sz/1024**3:.1f} GB",
                                 font=("Helvetica", 9), fg=C["muted"],
                                 bg=C["row"]).pack(side="left")

                    def make_rm(n):
                        def rm():
                            if not messagebox.askyesno(
                                    "Uninstall", f"Remove {n}?\nThis frees up disk space."):
                                return
                            def do_rm():
                                try:
                                    r = requests.delete(f"{OLLAMA_BASE}/api/delete",
                                                        json={"name": n}, timeout=10)
                                    def af():
                                        t3_stat.config(
                                            text=f"✓  {n} removed." if r.ok
                                                 else f"Failed to remove {n}.",
                                            fg=C["green"] if r.ok else C["red"])
                                        refresh_t3()
                                    root.after(0, af)
                                except Exception as ex:
                                    root.after(0, lambda: t3_stat.config(
                                        text=str(ex), fg=C["red"]))
                            Thread(target=do_rm, daemon=True).start()
                        return rm

                    tk.Button(row, text="Uninstall", font=("Helvetica", 9),
                              relief="flat", bd=0, bg=C["red_bg"], fg=C["red"],
                              activebackground="#fecaca", padx=8, pady=3,
                              cursor="hand2", command=make_rm(name)).pack(side="right")
            root.after(0, upd)
        Thread(target=do, daemon=True).start()

    t3_ref.config(command=refresh_t3)

    # Footer: copyright label at the bottom of the window
    ftr = tk.Frame(root, bg=C["bg"])
    ftr.pack(fill="x", side="bottom")
    tk.Frame(ftr, bg=C["border"], height=1).pack(fill="x")
    tk.Label(ftr,
             text="© 2026 Siva Shanmugam Mariappan and Ashwin Malshe. Licensed under the MIT License.",
             font=("Helvetica", 9), fg=C["muted"], bg=C["bg"]).pack(pady=6)

    # Startup + window close handler: trigger initial Ollama check and wire WM_DELETE_WINDOW
    Thread(target=do_check, daemon=True).start()
    root.after(800, refresh_t3)

    def on_close():
        _quit_ollama()
        server.should_exit = True
        root.destroy()
        sys.exit(0)

    root.protocol("WM_DELETE_WINDOW", on_close)
    root.mainloop()
